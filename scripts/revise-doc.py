"""
修订《技术方案.docx》，使其与当前代码仓库状态同步。
输出：技术方案-修订版.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy
import re

SRC = 'D:/game_hosted/hyacintech-stem-platform/技术方案-原始备份.docx'
DST = 'D:/game_hosted/hyacintech-stem-platform/技术方案-修订版.docx'

doc = Document(SRC)

# ============================================================
# Helper utilities
# ============================================================

def find_para_containing(text_fragment, start=0):
    """Return (index, paragraph) of first paragraph containing the given text."""
    for i, p in enumerate(doc.paragraphs):
        if i >= start and text_fragment in p.text:
            return i, p
    return None, None

def find_para_exact(text, start=0):
    """Return (index, paragraph) of first paragraph whose text exactly equals the given text."""
    for i, p in enumerate(doc.paragraphs):
        if i >= start and p.text.strip() == text.strip():
            return i, p
    return None, None

def find_table_containing(text_fragment, start=0):
    """Return (index, table) of first table containing the given text."""
    for i, t in enumerate(doc.tables):
        if i >= start:
            full = '\n'.join(cell.text for row in t.rows for cell in row.cells)
            if text_fragment in full:
                return i, t
    return None, None

def set_cell_text(cell, text):
    """Replace all text in a table cell, preserving the first run's formatting."""
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ''
    first_para = cell.paragraphs[0]
    if first_para.runs:
        first_para.runs[0].text = text
    else:
        first_para.add_run(text)

def replace_para_text(para, new_text):
    """Replace all runs' text in a paragraph with new_text, keeping first run's formatting."""
    # Clear all runs
    for run in para.runs:
        run.text = ''
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)

def replace_in_para(para, old, new):
    """Replace substring old with new within a paragraph's runs."""
    full = para.text
    if old not in full:
        return False
    # Simple approach: accumulate all run text, replace, write back to first run
    new_full = full.replace(old, new)
    for run in para.runs:
        run.text = ''
    if para.runs:
        para.runs[0].text = new_full
    else:
        para.add_run(new_full)
    return True

def insert_paragraph_after(anchor_para, text, style=None):
    """Insert a new paragraph after anchor_para."""
    new_p = doc.add_paragraph(text)
    if style:
        new_p.style = doc.styles[style]
    # Move the new paragraph (which is at the end) to right after anchor_para
    anchor_para._element.addnext(new_p._element)
    return new_p

def insert_paragraphs_after(anchor_para, items):
    """
    Insert multiple paragraphs after anchor_para (in reverse order so they
    end up in correct order). items is a list of (text, style) tuples.
    """
    # Process in reverse so that inserting each one after anchor_para
    # preserves the correct order
    for text, style in reversed(items):
        insert_paragraph_after(anchor_para, text, style)

def find_heading(text, level=None):
    """Find a heading paragraph by text content."""
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == text.strip():
            if level is None or p.style.name.startswith(f'Heading {level}'):
                return i, p
    return None, None

# ============================================================
# Chapter 3 — 用户与场景
# ============================================================

print("=== Chapter 3: 用户与场景 ===")

# 3.2 核心使用流程 (paragraphs 37-38)
idx37, p37 = find_para_containing('核心互动流程如下')
if p37:
    insert_paragraph_after(p37,
        '前端交互方面，系统在启动时自动执行 /api/health 健康诊断，若检测到配置、网络或鉴权异常，则在对话区顶部展示黄色诊断横幅。'
        '对话过程中，根据 LLM 返回的 next_action_type 字段，前端自动切换交互模式：ask_choice 渲染选项按钮，'
        'confirmation 渲染确认/取消按钮对，text_input 和 info 保留自由文本输入。'
        '用户可对已发送消息进行重发（Resend）或就地编辑（Inline Edit）后重新提交，实现对话纠错。',
        'Normal (Web)')
    print("  [3.2] 新增前端交互描述 → OK")
else:
    print("  [3.2] ✗ 未找到目标段落")

# ============================================================
# Chapter 4 — 技术思路
# ============================================================

print("\n=== Chapter 4: 技术思路 ===")

# 4.1 整体技术方案 (paragraph 43)
idx43, p43 = find_para_containing('深度嵌入式技术方案')
if p43:
    old_text = p43.text
    suffix = ('当前版本已从 Mock 响应迁移至真实 LLM 集成，通过原生 fetch 调用 OpenAI 兼容 API'
              '（支持 OpenAI / DeepSeek 双 Provider 自动检测），并实现了 JSON 输出合规三层联防机制。')
    if suffix not in old_text:
        replace_para_text(p43, old_text + suffix)
    print("  [4.1] 补充 LLM 集成现状 → OK")
else:
    print("  [4.1] ✗ 未找到目标段落")

# 4.2 数据流 (paragraph 46)
idx46, p46 = find_para_containing('数据流向：用户前端')
if p46:
    replace_in_para(p46,
        '数据流向：用户前端 → /api/chat → 状态机校验 → 安全过滤 → 提示词工厂组装 → 大模型 API → 结构化 JSON 响应 → 前端渲染',
        '数据流向：用户前端 → /api/chat → 安全过滤（checkBlacklistedKeywords）→ 提示词工厂组装（getPromptForPhase + injectSafetyConstraints）→ LLM Provider（原生 fetch + OpenAI 兼容协议）→ JSON 解析器（三层提取策略 + 启发式降级）→ 错误分类器 → ChatResponse → 前端渲染')
    print("  [4.2] 更新数据流描述 → OK")
else:
    print("  [4.2] ✗ 未找到目标段落")

# 4.3.1 接口定义
# --- Delete ExploreState code block (paragraphs 51-57) ---
# Paragraph 50 is "核心接口定义（TypeScript，对应代码仓库 app/models/）："
# Paragraphs 51-57 are the HTML Preformatted code lines of ExploreState
idx50, p50 = find_para_containing('核心接口定义（TypeScript')
if p50:
    # Find the ExploreState code block and remove it
    # ExploreState starts around idx50+1
    for i in range(idx50 + 1, min(idx50 + 10, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        if 'interface ExploreState' in p.text:
            # Remove this and subsequent code lines until ChatRequest
            # First, find where the ChatRequest interface starts
            end_idx = i
            for j in range(i + 1, min(i + 15, len(doc.paragraphs))):
                if 'interface ChatRequest' in doc.paragraphs[j].text:
                    end_idx = j
                    break
            # Remove paragraphs from i to end_idx-1
            for k in range(end_idx - 1, i - 1, -1):
                doc.paragraphs[k]._element.getparent().remove(doc.paragraphs[k]._element)
            print(f"  [4.3.1] 删除 ExploreState 接口 ({end_idx - i} 段) → OK")
            break

    # --- Replace ChatRequest code block ---
    # Now find the ChatRequest paragraph
    for i, p in enumerate(doc.paragraphs):
        if 'interface ChatRequest' in p.text and 'currentPhase' in p.text:
            # This is a multi-line pre block; replace its content
            # Find the HTML Preformatted paragraphs that form ChatRequest
            # They should be consecutive
            # Replace with correct ChatRequest
            replace_para_text(p, 'interface ChatRequest {')
            # Find the next few paragraphs and replace them
            next_lines = [
                '  message: string;',
                '  phase: number;',
                '  history: Message[];',
                '}',
            ]
            j = i + 1
            for line in next_lines:
                if j < len(doc.paragraphs) and ('}' in doc.paragraphs[j].text or
                    'currentPhase' in doc.paragraphs[j].text or
                    'conversationId' in doc.paragraphs[j].text or
                    doc.paragraphs[j].style.name == 'HTML Preformatted'):
                    replace_para_text(doc.paragraphs[j], line)
                    j += 1
                else:
                    break
            print("  [4.3.1] ChatRequest 更新 → OK")
            break

    # --- Replace ChatResponse code block ---
    for i, p in enumerate(doc.paragraphs):
        if 'interface ChatResponse' in p.text:
            replace_para_text(p, 'interface ChatResponse {')
            chat_response_lines = [
                '  dialogue: string;',
                "  next_action_type: 'ask_choice' | 'text_input' | 'confirmation' | 'info';",
                '  options?: string[];',
                '  phase_complete: boolean;',
                '}',
            ]
            j = i + 1
            for line in chat_response_lines:
                if j < len(doc.paragraphs):
                    replace_para_text(doc.paragraphs[j], line)
                    j += 1
            print("  [4.3.1] ChatResponse 更新 → OK")

            # Add PhaseEnum + PhaseData after the ChatResponse block
            # j is now after the closing brace
            if j < len(doc.paragraphs):
                anchor = doc.paragraphs[j - 1]  # the '}' line
                new_blocks = [
                    ('', 'Normal (Web)'),
                    ('enum PhaseEnum {', 'HTML Preformatted'),
                    ('  TopicSelection = 1,   // 选题定向', 'HTML Preformatted'),
                    ('  PlanDesign = 2,       // 方案设计', 'HTML Preformatted'),
                    ('  Execution = 3,        // 过程执行', 'HTML Preformatted'),
                    ('  DataAnalysis = 4,     // 数据分析', 'HTML Preformatted'),
                    ('  ResultsFormation = 5, // 成果成型', 'HTML Preformatted'),
                    ('  Reflection = 6        // 结果反思', 'HTML Preformatted'),
                    ('}', 'HTML Preformatted'),
                    ('', 'Normal (Web)'),
                    ('interface PhaseData {', 'HTML Preformatted'),
                    ('  // 每个阶段的特定数据结构', 'HTML Preformatted'),
                    ('  [PhaseEnum.TopicSelection]?: {', 'HTML Preformatted'),
                    ('    interest?: string;', 'HTML Preformatted'),
                    ('    selectedTopic?: string;', 'HTML Preformatted'),
                    ('    researchQuestion?: string;', 'HTML Preformatted'),
                    ('  };', 'HTML Preformatted'),
                    ('  [PhaseEnum.PlanDesign]?: {', 'HTML Preformatted'),
                    ('    variables?: { independent?: string; dependent?: string; control?: string[] };', 'HTML Preformatted'),
                    ('    materials?: string[];', 'HTML Preformatted'),
                    ('    procedure?: string[];', 'HTML Preformatted'),
                    ('  };', 'HTML Preformatted'),
                    ('  [PhaseEnum.Execution]?: { rawData?: any; observations?: string[] };', 'HTML Preformatted'),
                    ('  [PhaseEnum.DataAnalysis]?: { analyzedData?: any; findings?: string[] };', 'HTML Preformatted'),
                    ('  [PhaseEnum.ResultsFormation]?: { conclusion?: string; report?: string };', 'HTML Preformatted'),
                    ('  [PhaseEnum.Reflection]?: { improvements?: string[]; nextSteps?: string[] };', 'HTML Preformatted'),
                    ('}', 'HTML Preformatted'),
                    ('', 'Normal (Web)'),
                ]
                insert_paragraphs_after(anchor, new_blocks)
                print("  [4.3.1] PhaseEnum + PhaseData 追加 → OK")

            # Also add Message interface before the code blocks section if not present
            # Insert after paragraph 50 (the intro sentence)
            # Actually, the intro has already shifted. Let's find it again.
            for i2, p2 in enumerate(doc.paragraphs):
                if '核心接口定义（TypeScript' in p2.text:
                    msg_blocks = [
                        ('', 'Normal (Web)'),
                        ('// Message 类型已扩展，支持前端直接渲染交互组件', 'HTML Preformatted'),
                        ('interface Message {', 'HTML Preformatted'),
                        ('  id: string;', 'HTML Preformatted'),
                        ("  role: 'user' | 'assistant';", 'HTML Preformatted'),
                        ('  content: string;', 'HTML Preformatted'),
                        ('  options?: string[];', 'HTML Preformatted'),
                        ("  actionType?: 'ask_choice' | 'text_input' | 'confirmation' | 'info';", 'HTML Preformatted'),
                        ('  phaseComplete?: boolean;', 'HTML Preformatted'),
                        ("  status?: 'sending' | 'sent' | 'error';", 'HTML Preformatted'),
                        ('}', 'HTML Preformatted'),
                    ]
                    insert_paragraphs_after(p2, list(reversed(msg_blocks)))
                    print("  [4.3.1] Message 接口新增 → OK")
                    break
            break

# 4.3.2 提示词工厂
idx87, p87 = find_para_containing('"next_action_type": "ask_choice | ask_input | confirm"')
if p87:
    replace_in_para(p87,
        '"next_action_type": "ask_choice | ask_input | confirm",',
        '"next_action_type": "ask_choice | text_input | confirmation | info",')
    print("  [4.3.2] next_action_type 枚举值修正 → OK")

    # Add note after the prompt example block
    # Find the "【安全约束】" paragraph in the example
    idx_safety, p_safety = find_para_containing('【安全约束】')
    if p_safety:
        # Find the end of the example block (look for next heading or the paragraph after this block)
        note_text = (
            '注：以上为简化示意。当前版本的提示词已引入"输出格式（必须严格遵守）"声明和'
            '"严格禁止"清单——要求 LLM 的整个回复必须是纯 JSON 对象，不得有任何额外文字、'
            '引言或代码块包裹。完整提示词见仓库 app/prompts/ 目录。'
        )
        insert_paragraph_after(p_safety, note_text, 'Normal (Web)')
        print("  [4.3.2] 提示词示例说明追加 → OK")

# Add JSON three-layer defense as a new paragraph after 4.3.2 section
# Find "4.3.3" heading
idx433, p433 = find_heading('4.3.3 双层安全规则过滤器', level=4)
if idx433 and p433:
    json_defense_items = [
        ('JSON 输出保障机制（三层联防）', 'Heading 4'),
        ('为保证 LLM 稳定输出符合 ChatResponse Schema 的结构化 JSON，系统实现三层联防：', 'Normal (Web)'),
        ('Layer 1 — Provider 层：API 请求传入 response_format: { type: "json_object" }，在协议层面强制 LLM 输出 JSON。', 'Normal (Web)'),
        ('Layer 2 — Prompt 层：全部 6 个阶段提示词重写输出约束，明确要求纯 JSON 输出并以"严格禁止"清单限制额外文字、引言、代码块包裹等行为。', 'Normal (Web)'),
        ('Layer 3 — Parser 层：实现三层 JSON 提取策略——直接 JSON.parse → 正则提取 markdown 代码块 → 花括号匹配截取。若全部失败，则启用 heuristicExtract() 启发式降级函数，从纯自然语言文本中自动匹配编号列表（如 1. 2. 3.）提取选项，并通过确认性关键词（确认、确定、准备好）推断 action type。确保即使 LLM 完全忽略 JSON 指令，用户也能看到对话内容。', 'Normal (Web)'),
        ('自动重试：若首次以 response_format JSON 模式调用后解析失败，系统自动以无格式限制 + 注入显式 JSON 指令的方式重试一次（callLLM 函数内建逻辑，见 app/api/chat/route.ts）。', 'Normal (Web)'),
    ]
    insert_paragraphs_after(p433, list(reversed(json_defense_items)))
    print("  [4.3.2] 三层联防新增 → OK")
else:
    print("  [4.3.2] ✗ 未找到 4.3.3 标题，跳过三层联防插入")

# 4.3.3 安全过滤器 — Table 13
idx_t13, t13 = find_table_containing('app/lib/safety/blacklist.json')
if t13:
    for row in t13.rows:
        for cell in row.cells:
            if 'app/lib/safety/blacklist.json' in cell.text:
                new_text = cell.text.replace(
                    'app/lib/safety/blacklist.json',
                    'app/prompts/index.ts（BLACKLIST_KEYWORDS 常量数组）'
                )
                set_cell_text(cell, new_text)
                print("  [4.3.3] Table 13 黑名单路径修正 → OK")
                break

# 4.3.4 上下文管理器 — 段落 100 + Table 14
idx100, p100 = find_para_containing('上下文管理器（防遗忘 & 防注入）')
if p100:
    note_ctx = ('（注：上下文截断与摘要压缩机制规划中，当前版本将完整历史发送至 LLM；'
                '注入检测通过 prompt 安全约束间接实现。）')
    if note_ctx not in p100.text:
        replace_para_text(p100, p100.text + note_ctx)
        print("  [4.3.4] 上下文管理器标注 → OK")

# Table 14
idx_t14, t14 = find_table_containing('保留最近 10 轮对话')
if t14:
    for row in t14.rows:
        for cell in row.cells:
            if '保留最近 10 轮对话' in cell.text or '强制注入' in cell.text or '注入检测' in cell.text:
                if '（规划中）' not in cell.text:
                    set_cell_text(cell, cell.text + '（规划中）')
    print("  [4.3.4] Table 14 标注 → OK")

# 4.3.5 教师后台 — 段落 101 + Table 10, 15
idx101, p101 = find_para_containing('教师后台审核系统')
if p101:
    note_teacher = ('（注：教师后台审核系统为规划功能，当前版本尚未实现。'
                    '学生端六阶段探究流程已完整可运行。）')
    if note_teacher not in p101.text:
        replace_para_text(p101, p101.text + note_teacher)
        print("  [4.3.5] 教师后台标注 → OK")

# Table 10 — 教师后台行
idx_t10, t10 = find_table_containing('教师后台')
if t10:
    for row in t10.rows:
        for cell in row.cells:
            if cell.text.strip() == '教师后台' or '教师后台' in cell.text:
                # Find the cell next to it (description column)
                row_cells = row.cells
                for c in row_cells:
                    if 'Next.js' in c.text or '同技术栈' in c.text:
                        if '（规划中）' not in c.text:
                            set_cell_text(c, c.text + '（规划中）')
                break
    print("  [4.3.5] Table 10 教师后台行标注 → OK")

    # Also fix: 安全规则过滤器 + 上下文管理器 + 大模型 API + 数据持久层
    for row in t10.rows:
        first_cell_text = row.cells[0].text.strip() if row.cells else ''
        desc_cell = row.cells[2] if len(row.cells) >= 3 else None
        if not desc_cell:
            continue
        if '安全规则过滤器' in first_cell_text:
            if '前端输入校验 + API 入口拦截 + AI 层隐性识别，三重防护' in desc_cell.text:
                new_t = desc_cell.text.replace(
                    '前端输入校验 + API 入口拦截 + AI 层隐性识别，三重防护',
                    '前端输入校验 + API 入口拦截（双层硬拦截）+ prompt 层语义约束'
                )
                set_cell_text(desc_cell, new_t)
                print("  [4.3.5] Table 10 安全过滤器行修正 → OK")
        if '上下文管理器' in first_cell_text:
            if '（规划中）' not in desc_cell.text:
                set_cell_text(desc_cell, desc_cell.text + '（规划中）')
                print("  [4.3.5] Table 10 上下文管理器行标注 → OK")
        if '大模型 API' in first_cell_text:
            if 'OpenAI SDK v6.39' in desc_cell.text:
                new_t = desc_cell.text.replace(
                    'OpenAI SDK v6.39 接入；主力 GPT-4o + 备份 DeepSeek-V3 / 通义千问',
                    '原生 fetch + OpenAI 兼容协议自建 Provider；支持 OpenAI / DeepSeek 双 Provider 自动检测，默认模型 deepseek-chat / gpt-4o'
                )
                set_cell_text(desc_cell, new_t)
                print("  [4.3.5] Table 10 大模型 API 行修正 → OK")
        if '数据持久层' in first_cell_text:
            if 'Vercel KV（会话）+ PostgreSQL（用户/报告，规划中）' in desc_cell.text:
                new_t = desc_cell.text.replace(
                    'Vercel KV（会话）+ PostgreSQL（用户/报告，规划中）',
                    '当前为前端内存状态，无持久化（规划中：Vercel KV 会话存储 + PostgreSQL 用户/报告）'
                )
                set_cell_text(desc_cell, new_t)
                print("  [4.3.5] Table 10 数据持久层行修正 → OK")

# Table 15 — 教师后台功能表
idx_t15, t15 = find_table_containing('实时查看每个学生当前所处阶段')
if t15:
    # Add （规划中） to the first description cell
    for row in t15.rows:
        for cell in row.cells:
            if '实时查看每个学生' in cell.text:
                if '（规划中）' not in cell.text:
                    set_cell_text(cell, cell.text + '（规划中）')
                break
    print("  [4.3.5] Table 15 标注 → OK")

# Table 12 — 提示词工厂补充
idx_t12, t12 = find_table_containing('强制 JSON 输出')
if t12:
    # Find the paragraph right after table 12 and add a supplement
    # We can't easily find "after table", so let's find the next heading after table 12
    pass  # Will handle with a note near 4.3.2

# Table 19 — 技术可行性
idx_t19, t19 = find_table_containing('GPT-4o 作为主力模型')
if t19:
    for row in t19.rows:
        for cell in row.cells:
            if 'GPT-4o 作为主力模型' in cell.text:
                new_t = cell.text.replace(
                    '选用 GPT-4o 作为主力模型，DeepSeek-V3 / 通义千问作为国产备份。'
                    '仓库已通过 openai SDK v6.39 接入 OpenAI 兼容协议，'
                    '强大的指令遵循能力与 JSON Mode 能够稳定支撑提示词工厂的契合度与状态机的精准跳转。',
                    '支持 OpenAI / DeepSeek 双 Provider，默认模型通过 LLM_MODEL 环境变量配置'
                    '（DeepSeek 默认 deepseek-chat，OpenAI 默认 gpt-4o）。'
                    '仓库通过原生 fetch + 自建 OpenAICompatibleProvider 类接入 OpenAI 兼容协议，'
                    '无需引入 SDK 依赖。强大的指令遵循能力与 JSON Mode 能够稳定支撑提示词工厂的契合度与状态机的精准跳转。'
                )
                set_cell_text(cell, new_t)
                print("  [4.3] Table 19 模型描述修正 → OK")
                break

# 4.3.6 + 4.3.7 — Add new subsections before 4.4 or at end of 4.3
# Find the last 4.3.x heading before Chapter 5 or 4.4
# Look for "5. 创新点" heading
idx5, p5 = find_heading('5. 创新点', level=2)
if not idx5:
    idx5, p5 = find_heading('5. 创新点')

if idx5 and p5:
    new_sections = [
        ('', 'Normal (Web)'),
        ('4.3.6 LLM 错误诊断系统', 'Heading 4'),
        ('系统实现了完整的错误分类与诊断链路：', 'Normal (Web)'),
        ('错误分类器（app/lib/llm/errors.ts）：将 LLM API 调用中的各类异常（网络故障、超时、'
         'API Key 无效、余额不足、访问被拒、模型不存在、上下文溢出、内容过滤、限流、服务器过载等）'
         '归类为 15 种标准错误码，每种均配有中文用户提示信息。前端错误展示直接使用分类后的中文消息。', 'Normal (Web)'),
        ('健康检查端点（GET /api/health）：分三步诊断——① 配置校验（检查 API Key 是否为占位符或缺失），'
         '② 连通性测试（调用 /models 端点验证网络可达），③ 鉴权 + 模型可用性测试（发送 max_tokens: 1 的 '
         '最小 chat 请求）。返回 healthy / degraded / unhealthy 三档状态及每项检查的详细信息。', 'Normal (Web)'),
        ('前端诊断横幅：ChatInterface 挂载时自动调用 /api/health，若非 healthy 则在对话区顶部展示黄色警告横幅，'
         '逐项列出配置、网络、鉴权问题，用户可手动关闭。', 'Normal (Web)'),
        ('LLMError 结构化异常类：Provider 层通过 throw new LLMError(code, detail, httpStatus) 抛出可分类的异常，'
         'classifyError() 函数在 API 路由层统一捕获并转换为结构化错误响应。', 'Normal (Web)'),
        ('', 'Normal (Web)'),
        ('4.3.7 消息重发与编辑', 'Heading 4'),
        ('前端 ChatInterface 支持对话纠错交互：', 'Normal (Web)'),
        ('重发（Resend）：用户可对任意已发送消息点击重发按钮，系统删除该消息及其后的所有对话记录，'
         '以原消息内容重新发起 /api/chat 请求。适用于用户觉得 AI 回复不满意、想换个角度重新提问的场景。', 'Normal (Web)'),
        ('编辑（Edit）：用户可对任意已发送消息点击编辑按钮，在原位置出现可编辑文本框（textarea），'
         '修改内容后点击发送或按 Enter 提交，效果等同于"删除该消息及之后内容 + 以新内容重新发送"。'
         '按 Escape 可取消编辑。', 'Normal (Web)'),
    ]
    idx_before = idx5 - 1
    # Walk back to find the last paragraph before the "5. 创新点" heading
    anchor_5 = doc.paragraphs[idx_before] if idx_before >= 0 else None
    if anchor_5:
        insert_paragraphs_after(anchor_5, new_sections)
        print("  [4.3.6/4.3.7] 新增两个小节 → OK")
    else:
        print("  [4.3.6/4.3.7] ✗ 无法定位插入点")
else:
    print("  [4.3.6/4.3.7] ✗ 未找到第5章标题")

# ============================================================
# Chapter 5 — 创新点
# ============================================================

print("\n=== Chapter 5: 创新点 ===")

# Table 17 — 技术创新点 ④ (强制 JSON Schema)
idx_t17, t17 = find_table_containing('Function Calling')
if t17:
    for row in t17.rows:
        for cell in row.cells:
            if 'Function Calling' in cell.text and 'JSON Mode' in cell.text:
                new_t = cell.text.replace(
                    '通过 OpenAI Function Calling / JSON Mode 强制大模型输出符合预定义 Schema 的结构化数据，'
                    '配合前端组件实现"对话 + 选项按钮 + 阶段进度条"的丰富交互。',
                    '通过 response_format: { type: "json_object" } 参数 + 三层 JSON 解析策略（直接解析/代码块提取/花括号匹配）'
                    '+ 启发式自然语言降级兜底，确保即使 LLM 完全忽略 JSON 指令，前端也能从纯文本中提取对话内容和交互选项。'
                    '配合自动重试机制（JSON 模式失败后以无格式限制 + 显式 JSON 指令重试），实现最高可用性的结构化输出保障。'
                )
                set_cell_text(cell, new_t)
                print("  [5] Table 17 创新点④修正 → OK")
                break

# ============================================================
# Chapter 6 — 可行性
# ============================================================

print("\n=== Chapter 6: 可行性 ===")

# Table 20 — 已完成模块清单
idx_t20, t20 = find_table_containing('package.json、tsconfig.json')
if t20:
    # Remove rows referencing phantom files
    rows_to_check = []
    for i, row in enumerate(t20.rows):
        full_text = ' '.join(cell.text for cell in row.cells)
        if 'DEEPSEEK_SETUP.md' in full_text or 'TROUBLESHOOTING.md' in full_text:
            rows_to_check.append(i)
    for i in reversed(rows_to_check):
        tr = t20.rows[i]._tr
        t20._tbl.remove(tr)
        print(f"  [6.2] Table 20 删除幻影文件行 → OK")

    # Update OpenAI SDK row
    for row in t20.rows:
        for cell in row.cells:
            if 'OpenAI SDK 集成（v6.39' in cell.text:
                set_cell_text(cell, 'LLM Provider 集成（原生 fetch + OpenAI 兼容协议，支持 OpenAI / DeepSeek 双 Provider 自动检测）')
                print("  [6.2] Table 20 SDK描述修正 → OK")
                break

    # Update 安全过滤器 row
    for row in t20.rows:
        for cell in row.cells:
            if '双层安全过滤器组件雏形' in cell.text:
                set_cell_text(cell, '双层安全过滤器（前端 SafetyFilter + 后端 checkBlacklistedKeywords）')
                print("  [6.2] Table 20 安全过滤器描述更新 → OK")
                break

    # Add new rows for new modules
    # Find the last row and add after it
    new_modules = [
        ['LLM 错误分类系统（15 种错误码 + 中文提示）', 'app/lib/llm/errors.ts'],
        ['JSON 解析器（三层提取 + 启发式降级）', 'app/lib/llm/parser.ts'],
        ['LLM Provider 抽象层（OpenAI 兼容协议）', 'app/lib/llm/provider.ts'],
        ['健康检查 API 端点', 'app/api/health/route.ts'],
        ['消息重发/编辑前端交互', 'app/components/ChatInterface.tsx'],
    ]
    for mod_name, mod_path in new_modules:
        row = t20.add_row()
        row.cells[0].text = mod_name
        row.cells[1].text = mod_path
    print("  [6.2] Table 20 新增 5 行已完成模块 → OK")

# Table 23 — Roadmap M1-M3
idx_t23, t23 = find_table_containing('MVP 开发期')
if t23:
    for row in t23.rows:
        for cell in row.cells:
            if '接入真实 LLM 替换当前模拟响应' in cell.text:
                if '✓' not in cell.text and '已完成' not in cell.text:
                    set_cell_text(cell, cell.text + ' ✓（已完成）')
                    print("  [6.5] Table 23 LLM接入标记已完成 → OK")
            if '完成六阶段状态机、双层安全过滤器' in cell.text:
                if '✓' not in cell.text and '已完成' not in cell.text:
                    set_cell_text(cell, cell.text + ' ✓（已完成）')
                    print("  [6.5] Table 23 状态机+安全过滤器标记已完成 → OK")
            if '选题定向 + 方案设计两个阶段' in cell.text:
                if '六阶段' not in cell.text:
                    new_t = cell.text.replace(
                        '完成选题定向 + 方案设计两个阶段的完整提示词与前端组件',
                        '六阶段提示词已全部就绪 ✓（已完成）'
                    )
                    set_cell_text(cell, new_t)
                    print("  [6.5] Table 23 六阶段提示词标记已完成 → OK")

# ============================================================
# Save
# ============================================================

doc.save(DST)
print(f"\n{'='*60}")
print(f"修订完成 → {DST}")
print(f"请用 Word 打开做最终格式审查。")
print(f"{'='*60}")
